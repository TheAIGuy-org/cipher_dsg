import re
from datetime import datetime
from pathlib import Path
from typing import Any

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from openpyxl import load_workbook

from config.settings import settings

OUTPUT_DIR = settings.PROJECT_ROOT / "data/docx_output"

# Bullet characters detected everywhere in the document
_BULLET_RE = r"[•‣◦⁃\-]"

# Rendering fields added by Adobe Extract — used to identify injected elements
_RENDERING_KEYS = {"Bounds", "ClipBounds", "Font", "HasClip", "Lang",
                   "ObjectID", "TextSize", "attributes"}


def clean_text(value: Any) -> str:
    """
    Sanitize a cell or paragraph value before writing to the Word document.

    Removes _x000D_ carriage return artifacts from Adobe Extract, normalizes
    whitespace, and strips XML-incompatible control characters that crash python-docx.

    Args:
        value: Any value — converted to str internally.

    Returns:
        Cleaned string, stripped of leading/trailing whitespace.
    """
    if value is None:
        return ""
    value = str(value)
    value = re.sub(r"_x000D_", "", value)
    value = value.replace("\r", "").replace("\n", " ")
    value = re.sub(r"\s+", " ", value)
    value = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]", "", value)
    return value.strip()


def _strip_bullet_char(text: str) -> str:
    """
    Remove the leading bullet character from text.

    Word's List Bullet style adds its own bullet marker — keeping the original
    would result in double bullets.

    Args:
        text: Raw bullet text that may start with •, ‣, ◦, ⁃, or -.

    Returns:
        Text with leading bullet character and whitespace removed.
    """
    return re.sub(rf"^{_BULLET_RE}\s*", "", text)


def _split_inline_bullets(text: str) -> list[tuple[str, bool]]:
    """
    Detect bullet lines and split embedded bullets within a paragraph.

    Handles two cases:
    - Whole line starts with a bullet char (including dash) → single bullet item.
    - Bullet chars embedded mid-string (e.g. LLM returns "Intro • Item1 • Item2")
      → split on bullet chars, mark each part after the first as a bullet.

    Args:
        text: Paragraph text that may contain bullet characters.

    Returns:
        List of (text, is_bullet) tuples. is_bullet=True means render as List Bullet style.
    """
    stripped = text.strip()
    if re.match(rf"^{_BULLET_RE}\s", stripped):
        return [(_strip_bullet_char(stripped), True)]

    parts = re.split(r"[•‣◦⁃]", text)
    if len(parts) == 1:
        return [(text, False)]

    result = []
    for i, part in enumerate(parts):
        part = part.strip()
        if not part:
            continue
        result.append((part, i > 0))
    return result


def _add_toc_line(doc: Document, label: str, page_num: str) -> None:
    """
    Add a single TOC entry with a right-aligned dot-leader tab stop.

    Uses raw OOXML because python-docx does not expose tab stop configuration.
    The tab character must be a <w:tab/> XML element — a \\t text character
    does not render as a dot leader in Word.

    Args:
        doc:      The Document object to append the paragraph to.
        label:    The TOC entry label (e.g. "2.2.1  Reference formula").
        page_num: The page number string (e.g. "2").
    """
    p   = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    for existing in pPr.findall(qn("w:tabs")):
        pPr.remove(existing)

    tabs_el = OxmlElement("w:tabs")
    tab_el  = OxmlElement("w:tab")
    tab_el.set(qn("w:val"),    "right")
    tab_el.set(qn("w:leader"), "dot")
    tab_el.set(qn("w:pos"),    "8640")   # ~15.2 cm in twips
    tabs_el.append(tab_el)
    pPr.append(tabs_el)

    p.add_run(label)

    r_tab = OxmlElement("w:r")
    r_tab.append(OxmlElement("w:tab"))
    p._p.append(r_tab)

    p.add_run(page_num)


def struct_to_docx(struct_dict: dict[str, Any],path:str, product_code: str) -> Path:
    """
    Convert a struct elements list into a Word document saved to the output folder.

    Reads table data from xlsx files in data/extracted/<stem>/tables/ and renders
    each element according to its Path type (heading, paragraph, bullet, table, TOC).

    Elements tagged with _injected=True (prose_after content added by the updater)
    are rendered during skip_mode so they appear after the table but before the
    original cell-text dump.

    Args:
        struct_dict:  Dict with an "elements" key containing the full element list.
        path: PDF filename stem used to locate extracted table xlsx files
                      (e.g. "face_day_cream_1614322").

    Returns:
        Path to the saved .docx file in output/<path>_<timestamp>.docx.
    """
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

    _matches = list(settings.DOSSIER_DIR.glob(f"*_{path}.pdf"))
    if _matches:
        extracted_dir = settings.PROJECT_ROOT / "data" / "extracted" / _matches[0].stem
    else:
        extracted_dir = next(
            (settings.PROJECT_ROOT / "data" / "extracted").glob(f"*{path}*"),
            settings.PROJECT_ROOT / "data" / "extracted" / path,
        )

    doc             = Document()
    skip_mode       = False
    pending_bullet  = False
    last_was_bullet = False
    elements        = struct_dict.get("elements", [])
    consumed        = set()   # indices already absorbed by peek-forward TOC logic

    # Pre-compute injected indices — elements tagged with _injected=True were added
    # by the updater (prose_after) and have no Adobe Extract rendering fields
    injected_indices = {i for i, el in enumerate(elements) if el.get("_injected")}

    for idx, element in enumerate(elements):
        if idx in consumed:
            continue

        path       = element.get("Path", "")
        text       = element.get("Text", "") or ""
        file_paths = element.get("filePaths")

        # ── 1. TABLE FROM XLSX ────────────────────────────────────────────────
        if file_paths:
            skip_mode       = True   # skip cell text elements that follow in the struct
            last_was_bullet = False
            xlsx_file       = extracted_dir / file_paths[0]

            if not xlsx_file.exists():
                print(f"❌ Missing table: {xlsx_file}")
                continue

            print(f"📊 Inserting table: {xlsx_file}")
            wb   = load_workbook(xlsx_file, data_only=True)
            ws   = wb.active
            rows = [r for r in ws.iter_rows(values_only=True)
                    if any(c is not None and str(c).strip() for c in r)]

            if rows:
                table       = doc.add_table(rows=len(rows), cols=len(rows[0]))
                table.style = "Table Grid"
                for i, row in enumerate(rows):
                    for j, cell in enumerate(row):
                        table.rows[i].cells[j].text = clean_text(cell)

            doc.add_paragraph("")
            continue

        # ── 2. SKIP TABLE CELL TEXT ───────────────────────────────────────────
        if skip_mode:
            if "/Table" in path and not file_paths:
                continue              # table cell element — stay in skip_mode

            elif any(t in path for t in ("H1", "H2", "H3")):
                skip_mode = False     # heading — stop skipping, fall through to render

            elif idx in injected_indices:
                # Injected prose_after element — render it, keep skip_mode on
                # so the original Adobe cell-text dump that follows stays skipped
                if text and text.strip():
                    doc.add_paragraph(clean_text(text))
                continue

            else:
                continue              # unknown post-table element — skip

        # ── 3. TOC ENTRIES ────────────────────────────────────────────────────
        if "TOC" in path:
            if text and text.strip():
                if "\t" in text:
                    # Adobe already joined label\tpage_number in one element
                    parts    = text.split("\t", 1)
                    label    = parts[0].rstrip(". ")
                    page_num = parts[1].strip() if len(parts) > 1 else None
                    if page_num:
                        _add_toc_line(doc, label, page_num)
                    elif label:
                        doc.add_paragraph(label)
                else:
                    # Peek forward to collect page number fragment from next element
                    entry_parts = [clean_text(text)]
                    page_num    = None
                    j = idx + 1
                    while j < len(elements):
                        nxt      = elements[j]
                        nxt_path = nxt.get("Path", "")
                        nxt_text = (nxt.get("Text") or "").strip()
                        if "TOC" not in nxt_path:
                            break
                        if re.match(r"^[0-9]+[.][0-9]+", nxt_text) or re.match(
                            r"^(Table|Figure)\s+[0-9]+", nxt_text, re.IGNORECASE
                        ):
                            break
                        if nxt_text:
                            if re.match(r"^[0-9]+$", nxt_text):
                                page_num = nxt_text
                            else:
                                entry_parts.append(clean_text(nxt_text))
                        consumed.add(j)
                        j += 1

                    label = "  ".join(entry_parts).rstrip(". ")
                    if page_num:
                        _add_toc_line(doc, label, page_num)
                    else:
                        doc.add_paragraph(label)
            continue

        # "Table of contents / tables / figures" title — bold Normal, not Heading 1
        if text and re.match(r"^Table of (contents|tables|figures)", text.strip(), re.IGNORECASE):
            from docx.shared import Pt
            p   = doc.add_paragraph()
            run = p.add_run(clean_text(text))
            run.bold      = True
            run.font.size = Pt(12)
            continue

        # ── 4. BULLET SYMBOL (Lbl) ────────────────────────────────────────────
        if "Lbl" in path:
            pending_bullet = True   # next element is the bullet body
            continue

        # ── 5. BULLET BODY (LBody after Lbl) ─────────────────────────────────
        if pending_bullet:
            clean = _strip_bullet_char(clean_text(text))
            if clean:
                doc.add_paragraph(clean, style="List Bullet")
                last_was_bullet = True
            pending_bullet = False
            continue

        # ── 6. ORPHAN LBODY (no preceding Lbl) ───────────────────────────────
        if "LBody" in path:
            clean = _strip_bullet_char(clean_text(text))
            if clean:
                doc.add_paragraph(clean, style="List Bullet")
                last_was_bullet = True
            continue

        # ── 7. NORMAL CONTENT ─────────────────────────────────────────────────
        if text and text.strip():
            clean = clean_text(text)

            # Insert blank line to visually separate bullet block from normal text
            if last_was_bullet and not any(b for _, b in _split_inline_bullets(clean)):
                doc.add_paragraph("")
                last_was_bullet = False

            if "H1" in path:
                doc.add_heading(clean, level=1)
            elif "H2" in path:
                doc.add_heading(clean, level=2)
            elif "H3" in path:
                doc.add_heading(clean, level=3)
            else:
                for chunk_text, is_bullet in _split_inline_bullets(clean):
                    if is_bullet:
                        doc.add_paragraph(chunk_text, style="List Bullet")
                        last_was_bullet = True
                    else:
                        doc.add_paragraph(chunk_text)
        else:
            if last_was_bullet and not pending_bullet:
                last_was_bullet = False

    ts       = datetime.now().strftime("%Y%m%d_%H%M%S")
    out_path = OUTPUT_DIR / f"{product_code}_updated.docx"
    doc.save(out_path)
    print(f"\n✅ DOCX saved: {out_path}")
    return out_path